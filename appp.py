#!/usr/bin/env python
# coding: utf-8

# IMPORT LIBRARIES
import re
import pandas as pd
import streamlit as st
import docx
import en_core_web_sm
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import RegexpTokenizer
import pdfplumber
import PyPDF2
import nltk
import pickle as pk


nlp = en_core_web_sm.load()

background_image_url = "https://images.unsplash.com/photo-1477093782505-e10aaeb27c6d?q=80&w=1780&auto=format&fit=crop&ixlib=rb-4.0.3&ixid=M3wxMjA3fDB8MHxwaG90by1wYWdlfHx8fGVufDB8fHx8fA%3D%3D"

st.markdown(
    f"""
    <style>
    .stApp {{
        background-image: url("{background_image_url}");
        background-size: cover;
        background-attachment: fixed;
        background-repeat: no-repeat;
        background-position: center;
    }}
    </style>
    """,
    unsafe_allow_html=True
)

st.title('Resume Filtration')
st.markdown('<style>h1{color: Yellow;}</style>', unsafe_allow_html=True)
st.subheader('Welcome to the Resume Classification App')
st.markdown('<style>h2{color: white;}</style>', unsafe_allow_html=True)

# FUNCTIONS
def extract_skills(resume_text):
    nlp_text = nlp(resume_text)
    noun_chunks = nlp_text.noun_chunks
    tokens = [token.text for token in nlp_text if not token.is_stop]

    data = pd.read_csv(r"skills.csv")
    skills = list(data.columns.values)
    skillset = []

    for token in tokens:
        if token.lower() in skills:
            skillset.append(token)

    for token in noun_chunks:
        token = token.text.lower().strip()
        if token in skills:
            skillset.append(token)
    return [i.capitalize() for i in set([i.lower() for i in skillset])]

def getText(filename):
    fullText = ''
    if filename.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(filename)
        for para in doc.paragraphs:
            fullText += para.text + '\n'
    else:
        with pdfplumber.open(filename) as pdf_file:
            pdoc = PyPDF2.PdfFileReader(filename)
            number_of_pages = pdoc.getNumPages()
            page = pdoc.pages[0]
            page_content = page.extractText()
        for paragraph in page_content:
            fullText += paragraph
    return fullText

def display(doc_file):
    resume = []
    if doc_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(doc_file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        resume.append('\n'.join(full_text))
    else:
        with pdfplumber.open(doc_file) as pdf:
            pages = pdf.pages[0]
            resume.append(pages.extract_text())
    return resume

def preprocess(sentence):
    sentence = str(sentence)
    sentence = sentence.lower()
    sentence = sentence.replace('{html}', "")
    cleanr = re.compile('<.*?>')
    cleantext = re.sub(cleanr, '', sentence)
    rem_url = re.sub(r'http\S+', '', cleantext)
    rem_num = re.sub('[0-9]+', '', rem_url)
    tokenizer = RegexpTokenizer(r'\w+')
    tokens = tokenizer.tokenize(rem_num)
    filtered_words = [w for w in tokens if len(w) > 2 if not w in stopwords.words('english')]
    lemmatizer = WordNetLemmatizer()
    lemma_words = [lemmatizer.lemmatize(w) for w in filtered_words]
    return " ".join(lemma_words)

file_type = pd.DataFrame([], columns=['Uploaded File', 'Predicted Profile', 'Skills'])
filename = []
predicted = []
skills = []

model = pk.load(open(r'modelbagg.pkl', 'rb'))
Vectorizer = pk.load(open(r'vector.pkl', 'rb'))

upload_file = st.sidebar.file_uploader('Upload Resumes', type=['docx', 'pdf'], accept_multiple_files=True)

st.sidebar.header('Select Filter')
select = ['Peoplesoft Resume', 'SQL Developer', 'React Developer', 'Workday']
option = st.sidebar.selectbox('Fields', select)

for doc_file in upload_file:
    if doc_file is not None:
        filename.append(doc_file.name)
        cleaned = preprocess(display(doc_file))
        prediction = model.predict(Vectorizer.transform([cleaned]))[0]
        predicted.append(prediction)
        extText = getText(doc_file)
        skills.append(extract_skills(extText))

if len(predicted) > 0:
    file_type['Uploaded File'] = filename
    file_type['Predicted Profile'] = predicted
    file_type['Skills'] = skills

    st.subheader('Resume Analysis Results')
    st.write('### All Resumes')
    st.table(file_type.style.format())

    st.write(f'### {option}')
    filtered_df = file_type[file_type['Predicted Profile'] == option]
    st.table(filtered_df.style.format())

st.image("ab.jpeg", caption='Extensions in resume', width=700)
st.image("cd.jpeg", caption='Extensions in resume', width=700)
